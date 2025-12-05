# ==========================================
# 🧠 COSA FA (spiegato semplice)
# ------------------------------------------
# Questo programma crea un libro a partire dal tuo TOC:
# - carichi il TOC
# - controlli/correggi il TOC (anche con AI)
# - confermi
# - genera i testi (in blocchi da max 500 parole a chiamata)
# - scarichi DOCX/PDF con formattazione.
# ==========================================

# ==========================================
# 📦 IMPORT: prendo gli attrezzi che servono
# ==========================================
import io
import os
import re
import math
from dataclasses import dataclass, field
from typing import List, Optional

import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PyPDF2 import PdfReader

# PDF (per impaginare il libro)
from reportlab.lib.pagesizes import A4, letter  # (non usati in UI, ma safe to keep)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

# Rilevazione lingua (se presente, altrimenti user selects)
HAS_LANGID = False
try:
    import langid  # pip install langid
    HAS_LANGID = True
except Exception:
    HAS_LANGID = False

# OpenAI (per scrivere i testi veri, se c’è la chiave)
OPENAI_OK = False
try:
    from openai import OpenAI
    if os.getenv("OPENAI_API_KEY"):
        openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        OPENAI_OK = True
except Exception:
    OPENAI_OK = False

# ==========================================
# 🔧 COSTANTI (regole semplici e chiare)
# ==========================================
MAX_SUBGEN_WORDS = 500           # mai più di 500 parole per singola generazione
MIN_SECTION_WORDS_USEFUL = 250   # meno di così una sezione non “respira”
MAX_SECTION_WORDS_SOFT = 1500    # sopra questa soglia segnaliamo “troppo lunga”

# Mappa lingue “umane” → etichette per prompt
LANG_LABELS = {
    "auto": "auto",
    "it": "Italian",
    "en": "English",
    "es": "Spanish",
    "fr": "French",
}

# Formati pagina “da libro” (limitati a 6x9 e 8.5x11 come richiesto)
PAGE_SIZES = {
    "6x9": (6 * 72, 9 * 72),           # pollici → punti tipografici
    "8.5x11": (8.5 * 72, 11 * 72),
}

# Scelte font (DOCX + PDF verranno armonizzate più avanti)
FONT_CHOICES = ["Times New Roman", "Roboto", "Comfortaa"]

# Tono di voce (UI inglese usa lista dedicata)
TONE_CHOICES = ["Scientifico", "Colloquiale", "Narrativo"]

# ==========================================
# 🧱 SCATOLE DOVE METTO I DATI (modelli)
# ==========================================
@dataclass
class Section:
    title: str
    target_words: int = 0
    blocks: int = 0
    texts: List[str] = field(default_factory=list)  # testi finali dei sottoblocchi (concatenati a vista)

@dataclass
class Chapter:
    title: str
    sections: List[Section] = field(default_factory=list)
    target_words: int = 0
    blocks: int = 0

@dataclass
class BookPlan:
    title: str
    subtitle: str
    author: str
    total_words: int
    block_size: int
    chapters: List[Chapter] = field(default_factory=list)
    language_code: str = "auto"          # "auto", "it", "en", ...
    tone: str = "Conversational"         # UI inglese
    brief: str = ""                      # descrizione breve che guida lo stile
    pdf_page: str = "6x9"                # formato libro
    font_name: str = "Times New Roman"   # font preferito
    # parte associata a ciascun capitolo (stessa lunghezza di chapters)
    chapter_parts: List[Optional[str]] = field(default_factory=list)

# ==========================================
# 🖥️ IMPOSTAZIONI BASE DELLA PAGINA (UI in inglese)
# ==========================================
st.set_page_config(page_title="Book Agent - Book Generator", page_icon="📘", layout="wide")
st.title("📘 Book Agent - Book Generator")
st.caption("Upload TOC → review/approve → generate → download DOCX/PDF")

# Inizializzo “memoria” (serve dopo)
for key, default in {
    "chapters": None,
    "allocation_done": False,
    "detected_lang": "auto",
    "confirmed_toc_text": None,
    "generated_plan": None,
    "docx_bytes": None,
    "pdf_bytes": None,
}.items():
    if key not in st.session_state:
        st.session_state[key] = default


# ==========================================
# 📂 BLOCK 2 - TOC UPLOAD AND REVIEW (FINAL FIXED VERSION)
# ==========================================

st.subheader("📄 Step 1 - Upload or paste your TOC")

# ------------------------------------------
# Apply pending TOC (if any)
# ------------------------------------------
if "toc_text_editable" not in st.session_state:
    st.session_state["toc_text_editable"] = ""

if "detected_lang" not in st.session_state:
    st.session_state["detected_lang"] = "auto"

if "_last_uploaded_name" not in st.session_state:
    st.session_state["_last_uploaded_name"] = None


# ------------------------------------------
# File uploader
# ------------------------------------------
uploaded_file = st.file_uploader(
    "Upload TOC (DOCX, PDF or TXT) - optional",
    type=["docx", "pdf", "txt"]
)

def extract_toc_from_docx(file):
    doc = Document(file)
    lines = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt and not txt.isdigit() and len(txt) > 2:
            lines.append(txt)
    return "\n".join(lines)

def extract_toc_from_pdf(file):
    reader = PdfReader(file)
    out = []
    for page in reader.pages[:3]:
        txt = page.extract_text()
        if txt:
            for ln in txt.splitlines():
                ln = ln.strip()
                if 2 < len(ln) < 200:
                    out.append(ln)
    return "\n".join(out)

def extract_toc_from_txt(file):
    content = file.read().decode("utf-8", errors="ignore")
    lines = [ln for ln in content.splitlines() if ln.strip()]
    return "\n".join(lines)

def _chapter_word(lang_code: str) -> str:
    mapping = {"it": "Capitolo", "en": "Chapter", "es": "Capítulo", "fr": "Chapitre"}
    return mapping.get(lang_code, "Chapter")


# Se l'utente carica un file
if uploaded_file:
    fname = uploaded_file.name.lower()

    with st.spinner("Reading the TOC..."):
        if fname.endswith(".docx"):
            toc_text = extract_toc_from_docx(uploaded_file)
        elif fname.endswith(".pdf"):
            toc_text = extract_toc_from_pdf(uploaded_file)
        else:
            toc_text = extract_toc_from_txt(uploaded_file)

    detected = "auto"
    if toc_text.strip() and HAS_LANGID:
        try:
            detected = langid.classify(toc_text[:500])[0]
        except:
            detected = "auto"

    st.session_state["detected_lang"] = detected

    # Aggiorna text_edit SOLO se è un file nuovo
    if st.session_state["_last_uploaded_name"] != uploaded_file.name:
        st.session_state["toc_text_editable"] = toc_text
        st.session_state["_last_uploaded_name"] = uploaded_file.name

    st.success(f"Detected language: **{detected.upper()}**")


# ------------------------------------------
# TEXT AREA DEL TOC (stabile)
# ------------------------------------------

current_toc = st.session_state["toc_text_editable"]

new_toc = st.text_area(
    "Captured / pasted TOC:",
    value=current_toc,
    key="toc_input",
    height=330,
    help="Paste or edit your TOC here."
)

if new_toc != current_toc:
    st.session_state["toc_text_editable"] = new_toc
    current_toc = new_toc


# ------------------------------------------
# Bottoni
# ------------------------------------------
col1, col2 = st.columns(2)
with col1:
    confirm_toc = st.button("✅ Confirm this TOC")
with col2:
    refine_toc = st.button("🧠 Refine TOC with AI")


# ------------------------------------------
# AI REFINEMENT
# ------------------------------------------
if refine_toc:
    if not OPENAI_OK:
        st.error("OpenAI API key missing. Cannot refine TOC.")
    elif not current_toc.strip():
        st.error("TOC is empty. Upload or paste it first.")
    else:
        lang_code = st.session_state.get("detected_lang", "en")
        chap_word = _chapter_word(lang_code if lang_code in ["it", "en", "es", "fr"] else "en")

        with st.spinner("Refining TOC..."):
            prompt = (
                "You are a professional non-fiction book editor.\n"
                "Clean, normalize and structure the following TOC:\n"
                f"- Use '{chap_word} X' for chapters\n"
                "- Use numbering like 1.1, 1.2 for subsections\n"
                "- Keep all meaning\n"
                "- Improve clarity\n"
                "- Output ONLY the cleaned list, one line per heading\n\n"
                f"Original TOC:\n{current_toc}"
            )

            resp = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You refine book TOCs."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.4,
                max_tokens=800
            )
            refined = (resp.choices[0].message.content or "").strip()

        st.session_state["toc_text_editable"] = refined
        st.session_state["confirmed_toc_text"] = refined

        st.success("TOC refined.")
        st.rerun()


# ------------------------------------------
# Conferma TOC
# ------------------------------------------
if confirm_toc:
    if not current_toc.strip():
        st.error("TOC is empty. Paste or upload before confirming.")
    else:
        st.session_state["confirmed_toc_text"] = current_toc
        st.success("TOC confirmed. Proceed to Step 2.")


# ==========================================
# 🧮 BLOCK 3 - WORD ALLOCATION & 500-WORD BLOCKING
# ==========================================

# ---------- PARSING HELPERS ----------

def _strip_leading_markers(s: str) -> str:
    """
    Pulisce il titolo:
    - bullet iniziali
    - numeri (1., 1.1, 1.1.1, 1) ecc.
    - prefissi tipo 'SOTTOCAPITOLO', 'Subchapter', 'Section' + numero.
    """
    s = s.strip()
    # bullet
    s = re.sub(r"^[\-\*\u2022]+\s*", "", s)
    # numerazioni tipo 1., 1.1, 1.1.1
    s = re.sub(r"^\d+(?:\.\d+){0,3}[\.\)\-:]?\s*", "", s)
    # forme tipo ".1 "
    s = re.sub(r"^\.\d+\s*", "", s)
    # prefissi testuali di sottocapitolo / sezione
    s = re.sub(
        r"^(sottocapitolo|sottcapitolo|sotto capitolo|subchapter|sub-chapter|section)\s*\d*\s*[:\-\.\)]*\s*",
        "",
        s,
        flags=re.IGNORECASE,
    )
    return s.strip()

def _is_toc_label(ln: str) -> bool:
    """Riconosce TOC / INDEX / INDICE ecc. da ignorare come contenuto."""
    s = ln.strip().lower()
    return (
        s in {"toc", "t.o.c.", "index", "indice", "table of contents"} or
        "table of contents" in s
    )

def _is_part_label(ln: str) -> bool:
    """Riconosce PART / PARTE (livello strutturale, non generativo)."""
    s = ln.strip().lower()
    return bool(re.match(r"^(part|parte|partie)\b", s))

def _is_chapter_keyword_line(ln: str) -> bool:
    """
    Riconosce CAPITOLO / CHAPTER solo come parola intera all'inizio
    (non SOTTOCAPITOLO).
    """
    s = ln.strip().lower()
    if re.match(r"^(chapter|capitolo)\b", s):
        return True
    if re.match(r"^(chapter|capitolo)\s+\d+\b", s):
        return True
    return False

def _parse_allocation_from_title(raw: str):
    """
    Estrae titolo + parole/blocchi dalla fine della riga.

    Formati supportati:
    - '... [800]'
    - '... [800 words]'
    - '... [800 parole]'
    - '... (2 blocchi)'
    - '... (3 blocks)'

    Regola:
    - se contiene 'block|blocks|blocchi|blocco' → blocchi
    - altrimenti il numero è parole
    """
    text = raw.rstrip()
    words = 0
    blocks = 0

    m = re.search(r'[\[\(]([^\]\)]*)[\]\)]\s*$', text)
    if m:
        inner = m.group(1).strip()
        text = text[:m.start()].rstrip()
        num_match = re.search(r"(\d+)", inner)
        if num_match:
            n = int(num_match.group(1))
            if re.search(r"(block|blocks|blocchi|blocco)", inner, re.IGNORECASE):
                blocks = n
            else:
                words = n

    return text.strip(), words, blocks

def _looks_like_chapter_without_keyword(ln: str) -> bool:
    """
    Euristica per capitoli senza parola 'Chapter/Capitolo'.
    NON deve riconoscere sottocapitoli (contengono 'sott', 'sub', 'section').
    """
    s = ln.strip().lower()

    if "sott" in s or "sub" in s or "section" in s:
        return False

    if len(ln.split()) > 12:
        return False
    if not ln[:1].isalpha():
        return False
    if ln[:1].islower():
        return False
    if re.search(r"[.:;]$", ln):
        return False
    return True

def parse_confirmed_toc(toc_text: str):
    """
    Parsing del TOC confermato.

    Ritorna:
    - chapters: List[Chapter]
    - chapter_parts: List[Optional[str]] → titolo PARTE associato a ogni capitolo

    Regole:
    - INTRODUZIONE (senza sottocapitoli) diventa capitolo singolo.
    - PARTE X ... è solo struttura, non generativa.
    - CAPITOLO X Y → capitolo.
    - SOTTOCAPITOLO / subchapter / section → sottocapitolo foglia (unità generativa).
    """
    lines = [ln.rstrip() for ln in toc_text.splitlines() if ln.strip()]

    chapters: List[Chapter] = []
    chapter_parts: List[Optional[str]] = []

    current_chapter: Optional[Chapter] = None
    current_part: Optional[str] = None
    chap_idx = 0

    for raw in lines:
        ln = raw.strip()
        if not ln:
            continue

        # TOC / Indice
        if _is_toc_label(ln):
            continue

        # PARTE / PART
        if _is_part_label(ln):
            current_part = ln.strip()
            continue

        # CAPITOLO esplicito
        if _is_chapter_keyword_line(ln):
            chap_idx += 1
            title_clean, _, _ = _parse_allocation_from_title(ln)
            title_clean = _strip_leading_markers(title_clean)
            current_chapter = Chapter(title=title_clean)
            chapters.append(current_chapter)
            chapter_parts.append(current_part)
            continue

        # CAPITOLO euristico (es. "INTRODUZIONE", "PREFACE", ecc.)
        if current_chapter is None:
            if _looks_like_chapter_without_keyword(ln):
                chap_idx += 1
                title_clean, _, _ = _parse_allocation_from_title(ln)
                title_clean = _strip_leading_markers(title_clean)
                current_chapter = Chapter(title=title_clean)
                chapters.append(current_chapter)
                chapter_parts.append(current_part)
                continue
            else:
                # fallback: capitolo generico
                chap_idx += 1
                current_chapter = Chapter(title=f"Chapter {chap_idx}")
                chapters.append(current_chapter)
                chapter_parts.append(current_part)

        # Tutto il resto = sottocapitolo foglia
        title_clean, words, blocks = _parse_allocation_from_title(ln)
        title_clean = _strip_leading_markers(title_clean)

        if not title_clean:
            continue

        sec = Section(title=title_clean, target_words=words, blocks=blocks)
        current_chapter.sections.append(sec)

    # safety: se un capitolo non ha sottocapitoli espliciti,
    # il capitolo stesso è l'unità generativa
    for ch in chapters:
        if not ch.sections:
            ch.sections.append(Section(title=ch.title))

    return chapters, chapter_parts

def finalize_allocation_from_toc(chapters: List[Chapter]):
    """
    Normalizza parole/blocchi per ogni sezione foglia:

    - se blocks e non words → words = blocks * MAX_SUBGEN_WORDS
    - se words e non blocks → blocks = ceil(words / MAX_SUBGEN_WORDS)

    Ritorna:
    - chapters aggiornati
    - total_words
    - lista delle sezioni ancora senza allocazione
    """
    all_secs: List[Section] = [sec for ch in chapters for sec in ch.sections]

    for sec in all_secs:
        if sec.blocks and not sec.target_words:
            sec.target_words = sec.blocks * MAX_SUBGEN_WORDS
        elif sec.target_words and not sec.blocks:
            sec.blocks = max(1, math.ceil(sec.target_words / MAX_SUBGEN_WORDS))

    missing = [s for s in all_secs if s.target_words <= 0 or s.blocks <= 0]

    total_words = sum(sec.target_words for sec in all_secs)

    for ch in chapters:
        ch.target_words = sum(s.target_words for s in ch.sections)
        ch.blocks = sum(s.blocks for s in ch.sections)

    return chapters, total_words, missing

def rebuild_toc_from_plan(chapters: List[Chapter], chapter_parts: List[Optional[str]]) -> str:
    """
    Ricostruisce un TOC normalizzato nel formato:

    TOC
    INTRODUZIONE [X]          ← H1 singolo (senza parte e senza sottocapitoli reali)

    PARTE 1 I CAZZI           ← H1 contenitore
    CAPITOLO 1 CIAO           ← H2
    I CIAO CAZZI [1300]       ← H3 foglia
    I PEZZI [1500]            ← H3 foglia
    """
    lines = ["TOC"]
    last_part = None

    for ch, part in zip(chapters, chapter_parts):
        ch_title = ch.title.strip()

        # Capitolo "standalone" senza parte e con una sola sezione che ripete il titolo:
        # es: INTRODUZIONE, PREFACE, ecc.
        is_intro_like = (
            part is None and
            len(ch.sections) == 1 and
            ch.sections[0].title.strip().lower() == ch_title.lower()
        )

        if is_intro_like:
            sec = ch.sections[0]
            lines.append(f"{ch_title} [{sec.target_words}]")
            lines.append("")  # riga vuota per separare blocchi nel TOC
            continue

        # Usa PARTE come livello H1 contenitore, se presente
        if part and part != last_part:
            lines.append(part.strip())
            last_part = part

        # CAPITOLO (H2), senza allocazione (solo struttura)
        if ch_title:
            lines.append(ch_title)

        # SOTTOCAPITOLI: le vere foglie con allocazione
        for sec in ch.sections:
            sec_title = sec.title.strip()
            lines.append(f"{sec_title} [{sec.target_words}]")

    return "\n".join(lines).strip()


# ---------- UI STEP 2 ----------

st.subheader("🧩 Step 2 - Book data & allocation")

confirmed_toc = st.session_state.get("confirmed_toc_text", "") or ""

if not confirmed_toc.strip():
    st.warning("Please confirm the TOC in Step 1 first.")
else:
    # 1) Parsing TOC in struttura (PARTE → CAPITOLO → sottocapitoli)
    temp_chapters, chapter_parts = parse_confirmed_toc(confirmed_toc)

    all_secs = [sec for ch in temp_chapters for sec in ch.sections]
    missing_initial = [s for s in all_secs if s.target_words <= 0 and s.blocks <= 0]
    needs_per_section_input = len(missing_initial) > 0

    # 2) Parametri generali di generazione
    lang_code = st.selectbox(
        "Generation language",
        ["auto", "it", "en", "es", "fr"],
        index=["auto", "it", "en", "es", "fr"].index(
            st.session_state.get("detected_lang", "auto")
            if st.session_state.get("detected_lang", "auto") in ["it", "en", "es", "fr"]
            else "auto"
        ),
        help="Leave 'auto' to use the detected language from the TOC, or force a specific language."
    )

    TONE_CHOICES_EN = ["Scientific", "Conversational", "Narrative"]
    tone = st.selectbox("Tone of voice", TONE_CHOICES_EN, index=TONE_CHOICES_EN.index("Conversational"))

    brief = st.text_area(
        "Brief (what the model should optimize for)",
        height=120,
        placeholder="Example: beginner audience, practical style, real examples, avoid heavy jargon..."
    )

    pdf_page = st.selectbox("Page size", ["6x9", "8.5x11"], index=0)
    font_name = st.selectbox("Primary font", FONT_CHOICES, index=0)

    # 3) Metadati libro + parole per sezioni senza allocazione
    missing_specs = []

    with st.form("book_info_form"):
        title = st.text_input("Book title")
        subtitle = st.text_input("Subtitle")
        author = st.text_input("Author")

        if needs_per_section_input:
            st.warning(
                f"There are {len(missing_initial)} sections without allocation "
                f"(no words or blocks). Please specify target words for each."
            )
            for idx, sec in enumerate(missing_initial):
                v = st.number_input(
                    f"Words for section: '{sec.title}'",
                    min_value=MIN_SECTION_WORDS_USEFUL,
                    step=100,
                    value=MIN_SECTION_WORDS_USEFUL,
                    key=f"missing_words_{idx}",
                )
                missing_specs.append((sec, v))

        submitted_meta = st.form_submit_button("Save book data & compute allocation")

    if submitted_meta:
        # assegna parole alle sezioni mancanti
        for sec, v in missing_specs:
            sec.target_words = int(v)
            sec.blocks = max(1, math.ceil(sec.target_words / MAX_SUBGEN_WORDS))

        # 4) Finalizza allocazione
        chapters_alloc, total_words, missing_after = finalize_allocation_from_toc(temp_chapters)

        if missing_after:
            st.error("Some sections are still missing allocation. Check your TOC or per-section word counts.")
        else:
            # 5) Costruisce il piano libro
            plan_preview = BookPlan(
                title=title or "Title",
                subtitle=subtitle or "",
                author=author or "",
                total_words=total_words,
                block_size=MAX_SUBGEN_WORDS,
                chapters=chapters_alloc,
                language_code=lang_code,
                tone=tone,
                brief=brief.strip(),
                pdf_page=pdf_page,
                font_name=font_name,
                chapter_parts=chapter_parts,
            )

            st.session_state.generated_plan = plan_preview
            st.session_state.chapters = chapters_alloc
            st.session_state.allocation_done = True

            # 6) TOC normalizzato nello stato logico (verrà mostrato in Step 1 al prossimo rerun)
            new_toc = rebuild_toc_from_plan(chapters_alloc, chapter_parts)
            st.session_state["toc_text_editable"] = new_toc
            st.session_state["confirmed_toc_text"] = new_toc

            st.success(f"✅ Allocation ready. Total words: ~{total_words}. TOC has been normalized.")

 # 7) Preview allocazione con struttura PARTE → CAPITOLO → SOTTOCAPITOLO
if st.session_state.get("allocation_done") and st.session_state.get("generated_plan"):
    plan: BookPlan = st.session_state.generated_plan
    chapters_alloc = plan.chapters
    parts = plan.chapter_parts if plan.chapter_parts else [None] * len(chapters_alloc)

    st.markdown("### Allocation preview")

    last_part = None
    global_ch_idx = 0

    for ch, part in zip(chapters_alloc, parts):
        ch_title = ch.title.strip()

        # Capitolo H1 standalone (INTRODUZIONE, PREFACE, ecc.)
        is_intro_like = (
            part is None
            and len(ch.sections) == 1
            and ch.sections[0].title.strip().lower() == ch_title.lower()
        )

        if is_intro_like:
            global_ch_idx += 1
            sec = ch.sections[0]
            with st.expander(
                f"H1: {ch_title} — {sec.target_words} words — {sec.blocks} total blocks"
            ):
                st.write(
                    f"• This is a standalone H1 section (no parts/chapters below). "
                    f"Text will be generated directly at this level."
                )
            continue

        # Se cambia la PARTE, mostra il titolo della PARTE
        if part and part.strip() and part != last_part:
            st.markdown(f"**PART / PARTE:** {part.strip()}")
            last_part = part

        global_ch_idx += 1
        with st.expander(
            f"Chapter {global_ch_idx}: {ch_title} — {ch.target_words} words — {ch.blocks} total blocks"
        ):
            for j, sec in enumerate(ch.sections, start=1):
                st.write(
                    f"• Subchapter {j}: {sec.title} — {sec.target_words} words — "
                    f"{sec.blocks} blocks (≤{MAX_SUBGEN_WORDS} words each)"
                )

    st.info("When satisfied, proceed to Step 3: content generation.")



# ==========================================
# ✍️ BLOCK 4 - CONTENT GENERATION & EXPORT
# ==========================================

st.subheader("🖋️ Step 3 - Content generation & export")

# ---------- GENERATION HELPERS ----------

def _effective_language_label(plan: BookPlan) -> str:
    code = plan.language_code
    if code == "auto":
        det = st.session_state.get("detected_lang", "en")
        code = det if det in LANG_LABELS else "en"
    return LANG_LABELS.get(code, "English")

def _tone_instruction(tone: str) -> str:
    t = (tone or "").lower()
    if t.startswith("scien"):
        return "Use a precise, rigorous, evidence-based tone."
    if t.startswith("narr"):
        return "Use a narrative, evocative tone with smooth transitions."
    return "Use a clear, friendly, and practical tone."

def _generate_subchunk(prompt_sys: str, prompt_user: str) -> str:
    if not OPENAI_OK:
        return "[No API key configured.]"
    try:
        resp = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": prompt_sys},
                {"role": "user", "content": prompt_user},
            ],
            temperature=0.7,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        return f"[generation error] {e}"

def generate_block_text(
    plan: BookPlan,
    ch_title: str,
    sec_title: str,
    target_words: int,
    prev_summary: str = "",
    is_last_block: bool = False
) -> str:

    lang = _effective_language_label(plan)
    tone_ins = _tone_instruction(plan.tone)
    n_sub = max(1, math.ceil(target_words / MAX_SUBGEN_WORDS))
    words_per_sub = min(math.ceil(target_words / n_sub), MAX_SUBGEN_WORDS)

    sys = (
        "You are an expert non-fiction writer. "
        f"Write in {lang}. {tone_ins} Avoid repetition. "
        "Do not restate book/chapter/section titles. "
        "Write continuous prose (no lists unless necessary)."
    )

    final_chunks = []
    for idx in range(n_sub):
        note = "Start naturally." if idx == 0 else "Continue smoothly."
        if idx == n_sub - 1 and is_last_block:
            note += " Conclude naturally."

        context = []
        if plan.brief:
            context.append(f"Brief: {plan.brief}")
        if prev_summary:
            context.append(f"Previous context: {prev_summary}")

        user = (
            f"Book title: {plan.title}\nSubtitle: {plan.subtitle}\nAuthor: {plan.author}\n"
            f"Chapter: {ch_title}\nSection: {sec_title}\nTarget: ~{words_per_sub} words\n"
            f"{note}\n"
            + ("\n".join(context) if context else "")
        )

        txt = _generate_subchunk(sys, user)
        final_chunks.append(txt.strip())

    return " ".join(final_chunks).strip()

def generate_all_sections(plan: BookPlan):
    total_blocks = sum(sec.blocks for ch in plan.chapters for sec in ch.sections)
    if total_blocks <= 0:
        st.warning("No blocks to generate. Check your allocation.")
        return

    bar = st.progress(0, text="Writing in progress...")
    done = 0
    prev_summary = ""

    for ch in plan.chapters:
        for sec in ch.sections:
            sec.texts = []

            block_target = max(1, math.ceil(sec.target_words / max(1, sec.blocks)))

            for b in range(sec.blocks):
                text = generate_block_text(
                    plan,
                    ch_title=ch.title,
                    sec_title=sec.title,
                    target_words=block_target,
                    prev_summary=prev_summary,
                    is_last_block=(b == sec.blocks - 1),
                )
                sec.texts.append(text)

                words = re.split(r"\s+", text.strip())
                if len(words) > 120:
                    prev_summary = " ".join(words[:60]) + " ... " + " ".join(words[-40:])
                else:
                    prev_summary = text[:800]

                done += 1
                bar.progress(done / total_blocks, text=f"Blocks completed: {done}/{total_blocks}")

    bar.empty()
    st.success("✅ Content generation completed.")


# ---------- EXPORT HELPERS ----------

PDF_FONT_MAP = {
    "Times New Roman": "Times-Roman",
    "Roboto": "Helvetica",
    "Comfortaa": "Courier",
}

def _safe_filename(plan: BookPlan) -> str:
    base = f"{plan.title.strip()}_{plan.subtitle.strip()}" if plan.subtitle.strip() else plan.title.strip()
    base = re.sub(r"[^\w\-]+", "_", base).strip("_")
    return re.sub(r"_+", "_", base) or "book"

# TOC Word
def _add_docx_toc(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    p = doc.add_paragraph()
    run = p.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(_qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    p._p.append(fld)


# ---------- DOCX BUILDER ----------

def build_docx(plan: BookPlan, include_toc=True, include_copyright=False) -> bytes:
    doc = Document()

    sec = doc.sections[0]
    if plan.pdf_page == "8.5x11":
        from docx.shared import Cm
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.54)
    else:
        sec.page_width, sec.page_height = Inches(6), Inches(9)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(plan.title)
    r.bold = True
    r.font.size = Pt(26)

    if plan.subtitle.strip():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(plan.subtitle)
        r.font.size = Pt(16)

    for _ in range(12):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(plan.author)
    r.font.size = Pt(12)

    doc.add_page_break()

    # Copyright
    if include_copyright:
        p = doc.add_paragraph("Copyright & Disclaimer")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

        para = doc.add_paragraph(
            f"© {plan.author}. All rights reserved.\n\n"
            "No part of this publication may be reproduced, distributed, or transmitted in any form or by any means, "
            "including photocopying, recording, or other electronic or mechanical methods, without the prior written "
            "permission of the publisher, except in the case of brief quotations embodied in critical reviews.\n\n"
            "Disclaimer: The information in this book is provided for educational purposes only and does not constitute "
            "professional advice. Always consult a qualified professional for your specific situation."
        )
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_page_break()

    # TOC
    if include_toc:
        _add_docx_toc(doc)
        doc.add_page_break()

    # Font
    style = doc.styles["Normal"]
    style.font.name = plan.font_name
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), plan.font_name)
    except Exception:
        pass

    chapters = plan.chapters
    parts = plan.chapter_parts if plan.chapter_parts else [None] * len(chapters)
    last_part = None

    # Main content
    for ch, part in zip(chapters, parts):
        ch_title = ch.title.strip()

        # INTRODUZIONE / PREFACE (H1 diretto)
        is_intro_like = (
            part is None
            and len(ch.sections) == 1
            and ch.sections[0].title.strip().lower() == ch_title.lower()
        )

        if is_intro_like:
            p = doc.add_paragraph(ch_title)
            p.style = doc.styles["Heading 1"]

            sec_leaf = ch.sections[0]
            for text in sec_leaf.texts:
                para = doc.add_paragraph(text)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_page_break()
            continue

        # Parte (H1)
        if part and part != last_part:
            p = doc.add_paragraph(part.strip())
            p.style = doc.styles["Heading 1"]
            last_part = part

        # Capitolo (H2)
        if ch_title:
            p = doc.add_paragraph(ch_title)
            p.style = doc.styles["Heading 2"]

        # Sottocapitoli (H3)
        for sec_leaf in ch.sections:
            sec_title = sec_leaf.title.strip()
            if sec_title.lower() != ch_title.lower():
                p = doc.add_paragraph(sec_title)
                p.style = doc.styles["Heading 3"]

            for text in sec_leaf.texts:
                para = doc.add_paragraph(text)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ---------- PDF BUILDER ----------

class _TocDocTemplate(SimpleDocTemplate):
    def afterFlowable(self, f):
        if isinstance(f, Paragraph):
            nm = getattr(f.style, "name", "")
            if nm in ("H1", "H2", "H3"):
                level = {"H1": 0, "H2": 1, "H3": 2}[nm]
                self.notify("TOCEntry", (level, f.getPlainText(), self.canv.getPageNumber()))

def build_pdf(plan: BookPlan, include_toc=True, include_copyright=False) -> bytes:
    pagesize = PAGE_SIZES.get(plan.pdf_page, PAGE_SIZES["6x9"])
    buf = io.BytesIO()

    m = 2.54 * cm if plan.pdf_page == "8.5x11" else 2 * cm
    doc = _TocDocTemplate(buf, pagesize=pagesize, leftMargin=m, rightMargin=m, topMargin=m, bottomMargin=m)

    styles = getSampleStyleSheet()
    fnt = PDF_FONT_MAP.get(plan.font_name, "Times-Roman")

    H1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName=fnt, alignment=TA_LEFT, spaceBefore=14, spaceAfter=8)
    H2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName=fnt, alignment=TA_LEFT, spaceBefore=10, spaceAfter=6)
    H3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName=fnt, alignment=TA_LEFT, spaceBefore=8, spaceAfter=4)
    Body = ParagraphStyle("Body", parent=styles["BodyText"], fontName=fnt, alignment=TA_JUSTIFY, leading=14)
    TitleC = ParagraphStyle("TitleC", parent=styles["Title"], fontName=fnt, alignment=TA_CENTER, spaceAfter=20)
    SubC = ParagraphStyle("SubC", parent=styles["BodyText"], fontName=fnt, alignment=TA_CENTER, spaceBefore=6, spaceAfter=40)

    story = []

    # Title page
    story += [Spacer(1, 40), Paragraph(plan.title, TitleC)]
    if plan.subtitle.strip():
        story.append(Paragraph(plan.subtitle, SubC))
    story += [Spacer(1, pagesize[1] * 0.55), Paragraph(plan.author, SubC), PageBreak()]

    # Copyright
    if include_copyright:
        story += [
            Paragraph("Copyright & Disclaimer", H1),
            Paragraph(
                f"© {plan.author}. All rights reserved.<br/><br/>"
                "No part of this publication may be reproduced, distributed, or transmitted in any form or by any means, "
                "including photocopying, recording, or other electronic or mechanical methods, without the prior written "
                "permission of the publisher, except in the case of brief quotations embodied in critical reviews.<br/><br/>"
                "Disclaimer: The information in this book is provided for educational purposes only and does not constitute "
                "professional advice. Always consult a qualified professional for your specific situation.",
                Body,
            ),
            PageBreak(),
        ]

    # TOC
    if include_toc:
        toc = TableOfContents()
        toc.levelStyles = [
            ParagraphStyle(fontName=fnt, name="TOC1", leftIndent=20, firstLineIndent=-10, spaceBefore=6, leading=12),
            ParagraphStyle(fontName=fnt, name="TOC2", leftIndent=36, firstLineIndent=-10, spaceBefore=4, leading=12),
            ParagraphStyle(fontName=fnt, name="TOC3", leftIndent=52, firstLineIndent=-10, spaceBefore=2, leading=12),
        ]
        story += [Paragraph("Table of Contents", H1), Spacer(1, 12), toc, PageBreak()]

    chapters = plan.chapters
    parts = plan.chapter_parts if plan.chapter_parts else [None] * len(chapters)
    last_part = None

    for ch, part in zip(chapters, parts):
        ch_title = ch.title.strip()

        # INTRODUZIONE / PREFACE
        is_intro_like = (
            part is None
            and len(ch.sections) == 1
            and ch.sections[0].title.strip().lower() == ch_title.lower()
        )

        if is_intro_like:
            story.append(Paragraph(ch_title, H1))
            sec_leaf = ch.sections[0]
            for text in sec_leaf.texts:
                story.append(Paragraph(text.replace("\n", "<br/>"), Body))
                story.append(Spacer(1, 8))
            story.append(PageBreak())
            continue

        if part and part != last_part:
            story.append(Paragraph(part.strip(), H1))
            story.append(Spacer(1, 8))
            last_part = part

        story.append(Paragraph(ch_title, H2))
        story.append(Spacer(1, 4))

        for sec_leaf in ch.sections:
            sec_title = sec_leaf.title.strip()
            if sec_title.lower() != ch_title.lower():
                story.append(Paragraph(sec_title, H3))

            for text in sec_leaf.texts:
                story.append(Paragraph(text.replace("\n", "<br/>"), Body))
                story.append(Spacer(1, 6))

        story.append(PageBreak())

    doc.build(story)
    return buf.getvalue()


# ---------- UI BUTTONS ----------
if st.session_state.get("allocation_done") and st.session_state.get("generated_plan"):
    plan: BookPlan = st.session_state.generated_plan

    c1, c2 = st.columns(2)
    with c1:
        opt_toc = st.checkbox("Include Table of Contents", value=True)
    with c2:
        opt_copyright = st.checkbox("Include Copyright & Disclaimer page", value=False)

    if st.button("🚀 CONFIRM AND GENERATE CONTENT", type="primary", use_container_width=True):
        generate_all_sections(plan)
        try:
            st.session_state["docx_bytes"] = build_docx(plan, include_toc=opt_toc, include_copyright=opt_copyright)
            st.session_state["pdf_bytes"] = build_pdf(plan, include_toc=opt_toc, include_copyright=opt_copyright)
        except Exception as e:
            st.error(f"Export error: {e}")

    if st.session_state.get("docx_bytes") and st.session_state.get("pdf_bytes"):
        st.subheader("📥 Download your book")
        fname = _safe_filename(plan)

        c1, c2 = st.columns(2)
        with c1:
            st.download_button(
                "Download DOCX",
                data=st.session_state["docx_bytes"],
                file_name=f"{fname}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with c2:
            st.download_button(
                "Download PDF",
                data=st.session_state["pdf_bytes"],
                file_name=f"{fname}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
